from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from services import parte, parteAte
import os
import comtypes.client

from pdf2image.exceptions import (
 PDFInfoNotInstalledError,
 PDFPageCountError,
 PDFSyntaxError
)

def pdf_to_images(path, file):
	abs_path = os.path.join(path, file)
	poppler_abs_path = os.path.abspath(os.path.join('poppler-0.68.0', 'bin'))
	images = convert_from_path(abs_path, poppler_path = poppler_abs_path)
	name = parteAte(file, -1, '.')
	for i, image in enumerate(images):
		i+=1
		fname = name + str(i) + ".png"
		fpath = os.path.join(path, fname)
		image.save(fpath, "PNG")

def arquivosParaImagem(path_to_files):
	for root, dir_names, file_names in os.walk(path_to_files):
		if len(file_names) == 0:
			continue

		print(root)
		for file in file_names:
			if "pdf" in file:
				print(file)
				pdf_to_images(root, file)

		print('==================')

def gerar_documento(path_to_files, file_name, model_file = ''):
	level_root = path_to_files.count("\\")

	document = ''
	if model_file:
		document = Document(model_file)
	else:
		document = Document()

	for root, dir_names, file_names in os.walk(path_to_files):
		level_path = root.count("\\")
		level = level_path - level_root
		level +=1
		# header = parte(root,-1, "\\")
		# document.add_heading(header, level = level)

		if len(file_names) == 0:
			continue
		# level += 1
		for file in file_names:
			if "pdf" in file:
				 f_name = parteAte(file, -1, '.')
				 document.add_heading(f_name)
				 #document.add_page_break()
			if "png" in file:
				 abs_path = os.path.join(root, file)
				 document.add_picture(abs_path,width=Inches(7))
	file_name = file_name + '.docx'
	abs_file = os.path.join(path_to_files, file_name)
	document.save(abs_file)


def word_to_pdf(path, file_name):
	word_file = file_name + '.docx'
	pdf_file = file_name + '.pdf'
	wdFormatPDF = 17

	in_file = os.path.join(path, word_file)
	out_file = os.path.join(path, pdf_file)

	word = comtypes.client.CreateObject('Word.Application')
	doc = word.Documents.Open(in_file)
	doc.SaveAs(out_file, FileFormat=wdFormatPDF)
	doc.Close()
	word.Quit()


#percorreDiretorios()
# arquivosParaImagem()

# dirs = [ "DOC01","DOC02","DOC03","DOC04", "DOC05", "DOC06", "DOC07", "DOC08", "DOC09", "DOC10", "DOC11", "DOC12", "DOC13"]
# path_base = os.getcwd()
# for dir_path in dirs:
# 	path = "C:\\Users\\Lucas Salvador\\Desktop\\200715-LUCAS-RODAR-CODIGO\\" + dir_path
# 	print(path)
# 	gerar_documento(path, dir_path)
# 	word_to_pdf(path_base, dir_path)
# 	print("====")